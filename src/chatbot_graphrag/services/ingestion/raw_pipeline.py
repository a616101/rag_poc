"""
原始管線

處理沒有結構化元資料的原始文件（PDF, DOCX, HTML）的管線。

處理階段：
1. 讀取檔案內容
2. 抽取文字和元資料
3. 推斷文件類型
4. 分塊內容
5. 生成嵌入向量
6. 儲存到向量資料庫和圖譜
"""

import hashlib
import io
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from chatbot_graphrag.core.constants import DocType, DocStatus, PipelineType
from chatbot_graphrag.models.pydantic.ingestion import (
    Chunk,
    Document,
    DocumentMetadata,
    IngestJob,
    PipelineResult,
    PipelineStageResult,
)
from chatbot_graphrag.services.ingestion.chunkers import get_chunker

logger = logging.getLogger(__name__)


class RawPipeline:
    """
    處理原始文件（PDF, DOCX, HTML）的管線。

    處理階段：
    1. 讀取檔案內容
    2. 抽取文字和元資料
    3. 推斷文件類型
    4. 分塊內容
    5. 生成嵌入向量
    6. 儲存到向量資料庫和圖譜
    """

    # 支援的檔案副檔名
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".html", ".htm", ".txt", ".md"}

    # 副檔名到內容類型的映射
    CONTENT_TYPES = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".html": "text/html",
        ".htm": "text/html",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }

    def __init__(self):
        """初始化原始管線。"""
        self._initialized = True

    async def process_file(
        self,
        file_path: str | Path,
        job: IngestJob,
        enable_graph: bool = True,
    ) -> tuple[Document | None, list[Chunk], list[dict[str, Any]]]:
        """
        處理單一原始文件檔案。

        Args:
            file_path: 文件檔案路徑
            job: 父攝取任務
            enable_graph: 是否抽取圖譜實體

        Returns:
            (Document, Chunks 列表, 錯誤列表) 元組
        """
        file_path = Path(file_path)
        errors: list[dict[str, Any]] = []

        # 檢查檔案是否存在
        if not file_path.exists():
            errors.append({
                "stage": "file_read",
                "error": f"File not found: {file_path}",
            })
            return None, [], errors

        # 檢查副檔名是否支援
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            errors.append({
                "stage": "file_read",
                "error": f"Unsupported file type: {ext}",
            })
            return None, [], errors

        # 讀取檔案內容
        try:
            if ext in {".txt", ".md", ".html", ".htm"}:
                content = file_path.read_text(encoding="utf-8")
                raw_bytes = content.encode("utf-8")
            else:
                raw_bytes = file_path.read_bytes()
                content = await self._extract_text(raw_bytes, ext)
        except Exception as e:
            errors.append({
                "stage": "file_read",
                "error": str(e),
            })
            return None, [], errors

        return await self.process_content(
            content=content,
            raw_bytes=raw_bytes,
            filename=file_path.name,
            content_type=self.CONTENT_TYPES.get(ext, "application/octet-stream"),
            job=job,
            enable_graph=enable_graph,
        )

    async def process_content(
        self,
        content: str,
        raw_bytes: bytes,
        filename: str,
        content_type: str,
        job: IngestJob | None = None,
        enable_graph: bool = True,
    ) -> tuple[Document | None, list[Chunk], list[dict[str, Any]]]:
        """
        處理原始文件內容。

        Args:
            content: 抽取的文字內容
            raw_bytes: 原始檔案位元組
            filename: 原始檔名
            content_type: MIME 內容類型
            job: 父攝取任務
            enable_graph: 是否抽取圖譜實體

        Returns:
            (Document, Chunks 列表, 錯誤列表) 元組
        """
        errors: list[dict[str, Any]] = []
        stages: list[PipelineStageResult] = []

        # 階段 1：推斷元資料
        start_time = datetime.utcnow()
        try:
            metadata = await self._infer_metadata(content, filename)
            stages.append(PipelineStageResult(
                stage_name="infer_metadata",
                success=True,
                duration_ms=(datetime.utcnow() - start_time).total_seconds() * 1000,
            ))
        except Exception as e:
            errors.append({
                "stage": "infer_metadata",
                "error": str(e),
            })
            # 使用後備元資料
            metadata = DocumentMetadata(
                doc_type=DocType.GENERIC,
                title=Path(filename).stem,
            )

        # 階段 2：建立 Document
        start_time = datetime.utcnow()
        content_hash = hashlib.sha256(raw_bytes).hexdigest()
        doc_id = self._generate_doc_id(filename, content_hash)

        document = Document(
            id=doc_id,
            metadata=metadata,
            content_hash=content_hash,
            raw_content=content,
            pipeline=PipelineType.RAW,
            version=1,
        )

        stages.append(PipelineStageResult(
            stage_name="create_document",
            success=True,
            duration_ms=(datetime.utcnow() - start_time).total_seconds() * 1000,
        ))

        # 階段 3：分塊內容
        start_time = datetime.utcnow()
        try:
            chunker = get_chunker(metadata.doc_type)
            chunks = chunker.chunk(
                content=content,
                doc_id=doc_id,
                metadata=metadata,
                doc_version=1,
            )
            stages.append(PipelineStageResult(
                stage_name="chunk_content",
                success=True,
                output_count=len(chunks),
                duration_ms=(datetime.utcnow() - start_time).total_seconds() * 1000,
            ))
        except Exception as e:
            errors.append({
                "stage": "chunk_content",
                "error": str(e),
            })
            return document, [], errors

        logger.info(
            f"Processed raw document {doc_id}: {len(chunks)} chunks, "
            f"type={metadata.doc_type.value}"
        )

        return document, chunks, errors

    async def _extract_text(self, raw_bytes: bytes, ext: str) -> str:
        """
        從二進制文件內容中抽取文字。

        Args:
            raw_bytes: 文件位元組
            ext: 副檔名

        Returns:
            抽取的文字內容
        """
        if ext == ".pdf":
            return await self._extract_pdf_text(raw_bytes)
        elif ext in {".docx", ".doc"}:
            return await self._extract_docx_text(raw_bytes)
        else:
            # 嘗試解碼為文字
            try:
                return raw_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return raw_bytes.decode("utf-8", errors="ignore")

    async def _extract_pdf_text(self, raw_bytes: bytes) -> str:
        """從 PDF 位元組中抽取文字。"""
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, cannot extract PDF text")
            return ""
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""

    async def _extract_docx_text(self, raw_bytes: bytes) -> str:
        """從 DOCX 位元組中抽取文字。"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX text")
            return ""
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    async def _infer_metadata(
        self,
        content: str,
        filename: str,
    ) -> DocumentMetadata:
        """
        從內容和檔名推斷文件元資料。

        Args:
            content: 文件文字內容
            filename: 原始檔名

        Returns:
            推斷的 DocumentMetadata
        """
        title = Path(filename).stem
        doc_type = DocType.GENERIC

        # 嘗試從內容關鍵字推斷 doc_type
        content_lower = content.lower()

        if any(kw in content_lower for kw in ["步驟", "流程", "申請", "辦理"]):
            doc_type = DocType.PROCEDURE
        elif any(kw in content_lower for kw in ["醫師", "專長", "門診"]):
            doc_type = DocType.PHYSICIAN
        elif any(kw in content_lower for kw in ["位置", "樓層", "大樓"]):
            doc_type = DocType.GUIDE_LOCATION
        elif any(kw in content_lower for kw in ["交通", "停車", "公車", "捷運"]):
            doc_type = DocType.GUIDE_TRANSPORT
        elif any(kw in content_lower for kw in ["團隊", "科別", "服務"]):
            doc_type = DocType.HOSPITAL_TEAM
        elif any(kw in content_lower for kw in ["問答", "常見問題", "faq", "q&a"]):
            doc_type = DocType.FAQ

        # 嘗試從內容中抽取標題
        lines = content.split("\n")
        for line in lines[:5]:  # 檢查前 5 行
            line = line.strip()
            if line and len(line) < 100:
                # 檢查標題模式
                if line.startswith("#"):
                    title = line.lstrip("#").strip()
                    break
                elif not line.startswith(("http", "www", "//")):
                    title = line
                    break

        # 偵測語言
        language = "zh-TW"  # 預設
        if re.search(r"[a-zA-Z]{20,}", content):  # Long English sequence
            language = "en"

        return DocumentMetadata(
            doc_type=doc_type,
            title=title,
            language=language,
        )

    def _generate_doc_id(
        self,
        filename: str,
        content_hash: str,
    ) -> str:
        """生成確定性的文件 ID。"""
        name_slug = Path(filename).stem.lower().replace(" ", "_")[:20]
        hash_prefix = content_hash[:8]
        return f"d_raw_{name_slug}_{hash_prefix}"

    async def process_directory(
        self,
        directory: str | Path,
        job: IngestJob,
        extensions: Optional[set[str]] = None,
        enable_graph: bool = True,
    ) -> PipelineResult:
        """
        處理目錄中所有支援的檔案。

        Args:
            directory: 目錄路徑
            job: 父攝取任務
            extensions: 要處理的特定副檔名（預設為所有支援的）
            enable_graph: 是否抽取圖譜實體

        Returns:
            包含所有處理過的文件和 chunks 的 PipelineResult
        """
        directory = Path(directory)
        started_at = datetime.utcnow()

        all_documents: list[Document] = []
        all_chunks: list[Chunk] = []
        all_errors: list[dict[str, Any]] = []
        stages: list[PipelineStageResult] = []

        # 決定要處理的副檔名
        target_extensions = extensions or self.SUPPORTED_EXTENSIONS

        # 找到所有匹配的檔案
        files: list[Path] = []
        for ext in target_extensions:
            files.extend(directory.glob(f"**/*{ext}"))

        logger.info(f"Found {len(files)} files to process")

        for file_path in files:
            doc, chunks, errors = await self.process_file(
                file_path=file_path,
                job=job,
                enable_graph=enable_graph,
            )

            if doc:
                all_documents.append(doc)
                all_chunks.extend(chunks)

            if errors:
                for err in errors:
                    err["file"] = str(file_path)
                all_errors.extend(errors)

        stages.append(PipelineStageResult(
            stage_name="process_files",
            success=len(all_errors) == 0,
            input_count=len(files),
            output_count=len(all_documents),
            duration_ms=(datetime.utcnow() - started_at).total_seconds() * 1000,
        ))

        return PipelineResult(
            job_id=job.id,
            pipeline=PipelineType.RAW,
            success=len(all_errors) == 0,
            stages=stages,
            documents=all_documents,
            chunks=all_chunks,
            total_duration_ms=(datetime.utcnow() - started_at).total_seconds() * 1000,
            started_at=started_at,
            completed_at=datetime.utcnow(),
        )


# 單例實例
raw_pipeline = RawPipeline()
